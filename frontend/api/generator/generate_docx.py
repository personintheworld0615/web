from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io

def set_cell_background(cell, fill_color):
    """Helper to set background color of a table cell."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def build_docx_to_stream(data, stream):
    doc = Document()
    meta = data["meta"]
    
    # ── Title & Subtitle ──────────────────────────────────────────────────────
    title = doc.add_heading(meta["title"], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph(meta["subtitle"])
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    
    doc.add_paragraph("_" * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Description
    desc_lines = [l.strip() for l in meta["description"].split("\n") if l.strip()]
    for line in desc_lines:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if any(t in line.lower() for t in ["counselors may not require", "no one may add", "NOTE:"]):
            p.runs[0].bold = True
    
    doc.add_paragraph("_" * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Revision Note
    if meta.get("revision_note"):
        rev = doc.add_paragraph(meta["revision_note"])
        rev.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rev.runs[0].font.size = Pt(8)
        rev.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Participant Fields
    doc.add_paragraph() # Spacer
    table = doc.add_table(rows=2, cols=3)
    table.style = 'Table Grid'
    fields = meta["participant_fields"]
    
    for i in range(min(len(fields), 6)):
        row = i // 3
        col = i % 3
        cell = table.rows[row].cells[col]
        cell.text = f"{fields[i]['label']}: "
        cell.paragraphs[0].runs[0].bold = True

    doc.add_page_break()

    # ── Requirements ──────────────────────────────────────────────────────────
    for req in data["requirements"]:
        build_item(doc, req, level=0)

    # ── Appendices ────────────────────────────────────────────────────────────
    for app in data.get("appendices", []):
        doc.add_page_break()
        doc.add_heading(app["title"], level=1)
        
        if app["type"] == "budget_table":
            rows = app.get("rows", 10)
            cols = len(app.get("headers", ["Item", "Cost"]))
            t = doc.add_table(rows=rows + 1, cols=cols)
            t.style = 'Table Grid'
            # Headers
            for i, h in enumerate(app.get("headers", [])):
                t.rows[0].cells[i].text = h
                t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
                set_cell_background(t.rows[0].cells[i], "DDE3EF")

    doc.save(stream)

def build_item(doc, item, level):
    # Requirement Header
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2 * level)
    
    if item.get("id"):
        run = p.add_run(f"{item['id']}. ")
        run.bold = True
    
    p.add_run(item["text"])
    
    # Component Rendering
    comp = item.get("component")
    if comp == "checklist":
        for option in item.get("options", []):
            p = doc.add_paragraph(f"☐ {option}", style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.2 * (level + 1))
    
    elif comp == "tracking_grid":
        rows = item.get("rows", 5)
        cols = item.get("cols", 7)
        t = doc.add_table(rows=rows + 1, cols=cols + 1)
        t.style = 'Table Grid'
        # Labels
        for i, label in enumerate(item.get("col_labels", [])):
            if i + 1 < cols + 1:
                t.rows[0].cells[i+1].text = label
        for i, label in enumerate(item.get("row_labels", [])):
            if i + 1 < rows + 1:
                t.rows[i+1].cells[0].text = label

    elif comp == "answer_box":
        # In Word, we'll just add a few empty lines or a single cell table
        t = doc.add_table(rows=1, cols=1)
        t.style = 'Table Grid'
        cell = t.rows[0].cells[0]
        cell.height = Inches(item.get("lines", 3) * 0.25)
        set_cell_background(cell, "F5F5F5")
        doc.add_paragraph() # Spacer

    # Sub-requirements
    for sub in item.get("sub", []):
        build_item(doc, sub, level + 1)
